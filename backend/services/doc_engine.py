import os
import re
import shutil
import zipfile
import tempfile
import httpx
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from pathlib import Path
from datetime import date

OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL = os.getenv("OPENROUTER_MODEL", "openai/gpt-4.1-mini")

MOIS = ["janvier","février","mars","avril","mai","juin",
        "juillet","août","septembre","octobre","novembre","décembre"]

def _today_fr() -> str:
    d = date.today()
    return f"{d.day} {MOIS[d.month - 1]} {d.year}"


def _clean_title(title: str) -> str:
    title = re.sub(r'[\(\[]?\s*[HhXx]\s*[/\-]\s*[FfHh]\s*[\)\]]?', '', title).strip()
    if '/' in title:
        title = title.split('/')[0].strip()
    return title.strip(' -–')


def _slugify(text: str, max_len: int = 30) -> str:
    text = re.sub(
        r'\b(S\.?A\.?S(?:U)?|S\.?A(?!S)|SARL|SNC|S\.?E\.?S|N\.?V|B\.?V|Ltd|Inc\.?|Corp\.?)\b',
        '', text, flags=re.IGNORECASE
    )
    text = re.sub(r'[^a-zA-Z0-9\s]', '', text)
    text = ''.join(w.capitalize() for w in text.split())
    return text[:max_len] or "Entreprise"

def call_llm(prompt: str, max_tokens: int = 1500) -> tuple[str, int]:
    key = os.getenv("OPENROUTER_API_KEY")
    if not key:
        raise RuntimeError("OPENROUTER_API_KEY manquant dans .env")
    resp = httpx.post(
        OPENROUTER_URL,
        headers={
            "Authorization": f"Bearer {key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "http://localhost:5173",
            "X-Title": "applicator",
        },
        json={
            "model": MODEL,
            "max_tokens": max_tokens,
            "messages": [{"role": "user", "content": prompt}],
        },
        timeout=90,
    )
    resp.raise_for_status()
    data = resp.json()
    text = data["choices"][0]["message"]["content"].strip()
    tokens = data.get("usage", {}).get("total_tokens", 0)
    return text, tokens

def _extract_docx_text(docx_path: Path) -> str:
    from docx import Document
    doc = Document(str(docx_path))
    return "\n".join(p.text for p in doc.paragraphs)


def _extract_headline(text: str) -> str:
    result = []
    for line in text.split("\n"):
        if not line.strip() and result:
            break
        if line.strip():
            result.append(line.strip())
        if len(result) >= 6:
            break
    return "\n".join(result)

def _extract_candidate_name(lm_text: str) -> str:
    lines = [l.strip() for l in lm_text.split("\n") if l.strip()]
    for line in lines[:3]:
        words = line.split()
        if 2 <= len(words) <= 4 and not any(c.isdigit() or c == '@' for c in line):
            return line
    for line in reversed(lines):
        words = line.split()
        if 2 <= len(words) <= 4 and not any(c.isdigit() or c == '@' for c in line):
            return line
    return "Candidat"

def adapt_cv_headline(current_headline: str, job_title: str, job_company: str) -> tuple[str, int]:
    clean = _clean_title(job_title)

    lines = current_headline.split("\n")
    title_line = ""
    for line in lines:
        if line.strip() and not any(c.isdigit() for c in line):
            title_line = line.strip()
            break
    if not title_line and lines:
        title_line = lines[0].strip()

    new_headline = current_headline.replace(title_line, clean, 1)
    return new_headline, 0


def adapt_cover_letter(lm_text: str, job_title: str, job_company: str, job_location: str) -> tuple[str, int]:
    clean = _clean_title(job_title)
    today = _today_fr()
    prompt = f"""Tu dois adapter cette lettre de motivation pour une nouvelle candidature.

Lettre originale :
---
{lm_text}
---

Informations sur le nouveau poste :
- Entreprise : {job_company}
- Poste : {clean}
- Localisation du poste : {job_location}
- Date du jour : {today}

En lisant la lettre comme un humain, effectue ces modifications :
1. Dans l'en-tête destinataire : remplace le nom de l'entreprise après "À l'attention de" par "{job_company}". Laisse "Service Finance" ou toute autre ligne de service telle quelle.
2. Intitulé du poste dans l'objet et dans le corps → "{clean}"
3. Dans la ligne "Ville, le DATE" : remplace UNIQUEMENT la date par "{today}", laisse la ville telle quelle
4. Dans l'adresse du destinataire : remplace la ville/code postal par "{job_location}"
5. Dans le corps de la lettre : remplace chaque occurrence de l'ancien nom d'entreprise par "{job_company}"

Ne modifie rien d'autre. Retourne uniquement la lettre complète."""
    return call_llm(prompt, max_tokens=2000)


def adapt_email(lm_text: str, job_title: str, job_company: str) -> tuple[str, int]:
    clean = _clean_title(job_title)
    prompt = f"""Voici une lettre de motivation :

---
{lm_text[:600]}
...
{lm_text[-300:]}
---

Complète ce template d'email de candidature en remplaçant uniquement les balises [] :

Objet: Candidature – {clean} – [NOM PRENOM du candidat extrait de la lettre]

Bonjour,

Je vous adresse ma candidature pour le poste de {clean} au sein de {job_company}. Vous trouverez en pièces jointes mon CV ainsi que ma lettre de motivation.
Je reste à votre disposition pour toute information complémentaire.

Cordialement,

[NOM PRENOM]
[TELEPHONE trouvé dans la lettre — supprimer la ligne si absent]

Ne modifie rien d'autre. Retourne uniquement l'email complet (objet + corps)."""
    return call_llm(prompt, max_tokens=300)


def _patch_cv_title_in_docx(docx_path: Path, old_title: str, new_title: str) -> None:
    from docx import Document
    doc = Document(str(docx_path))
    for para in doc.paragraphs:
        if para.text.strip() == old_title:
            if para.runs:
                para.runs[0].text = new_title
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.text = new_title
            break
    doc.save(str(docx_path))


def _apply_text_to_docx(docx_path: Path, new_text: str) -> None:
    from docx import Document
    doc = Document(str(docx_path))
    new_lines = new_text.split("\n")
    new_idx = 0
    for para in doc.paragraphs:
        if new_idx >= len(new_lines):
            if para.runs:
                para.runs[0].text = ""
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.text = ""
            continue
        new_line = new_lines[new_idx]
        new_idx += 1
        if para.runs:
            para.runs[0].text = new_line
            for run in para.runs[1:]:
                run.text = ""
        else:
            para.text = new_line
    doc.save(str(docx_path))

def _parse_email_output(raw: str) -> tuple[str, str]:
    lines = raw.strip().split("\n")
    subject = "Candidature"
    body_start = 0
    for i, line in enumerate(lines):
        if line.lower().startswith("objet:") or line.lower().startswith("sujet:"):
            subject = line.split(":", 1)[1].strip()
            body_start = i + 1
            break
    while body_start < len(lines) and not lines[body_start].strip():
        body_start += 1
    body = "\n".join(lines[body_start:]).strip()
    return subject, body


def build_eml(subject: str, body: str, cv_path: Path, lm_path: Path) -> bytes:
    msg = MIMEMultipart()
    msg["Subject"] = subject
    msg["From"] = "moi@email.com"
    msg["To"] = "recruteur@entreprise.com"
    msg.attach(MIMEText(body, "plain", "utf-8"))
    for path in [cv_path, lm_path]:
        with open(path, "rb") as f:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f'attachment; filename="{path.name}"')
        msg.attach(part)
    return msg.as_bytes()


async def generate_documents(
    cv_bytes: bytes,
    cv_ext: str,
    lm_bytes: bytes,
    lm_ext: str,
    jobs: list[dict],
) -> tuple[str, int]:
    import asyncio

    if cv_ext not in (".docx", ".doc"):
        raise ValueError(f"CV : DOCX uniquement (reçu : {cv_ext}).")
    if lm_ext not in (".docx", ".doc"):
        raise ValueError(f"Lettre : DOCX uniquement (reçu : {lm_ext}).")

    loop = asyncio.get_event_loop()
    tmp_dir = Path(tempfile.mkdtemp())

    cv_template = tmp_dir / f"cv_template{cv_ext}"
    lm_template = tmp_dir / f"lm_template{lm_ext}"
    cv_template.write_bytes(cv_bytes)
    lm_template.write_bytes(lm_bytes)

    cv_text     = _extract_docx_text(cv_template)
    lm_text     = _extract_docx_text(lm_template)
    cv_headline = _extract_headline(cv_text)
    candidate   = _extract_candidate_name(lm_text)

    total_tokens = 0
    zip_name = _slugify(jobs[0].get("company", "candidature")) if len(jobs) == 1 else "candidatures"
    zip_path = tmp_dir / f"{zip_name}.zip"

    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for job in jobs:
            job_title    = job.get("title",       "Poste")      or "Poste"
            job_company  = job.get("company",     "Entreprise") or "Entreprise"
            job_location = job.get("location",    "")           or ""
            slug         = _slugify(job_company)

            (new_headline, t1), (new_lm, t2), (email_raw, t3) = await asyncio.gather(
                loop.run_in_executor(None, adapt_cv_headline,  cv_headline, job_title, job_company),
                loop.run_in_executor(None, adapt_cover_letter, lm_text,     job_title, job_company, job_location),
                loop.run_in_executor(None, adapt_email,        lm_text,     job_title, job_company),
            )
            total_tokens += t1 + t2 + t3

            cv_name = f"CV {candidate}.docx"
            lm_name = f"LM {candidate}.docx"

            cv_out = tmp_dir / cv_name
            shutil.copy2(cv_template, cv_out)
            new_cv_text = cv_text.replace(cv_headline, new_headline, 1)
            _apply_text_to_docx(cv_out, new_cv_text)

            lm_out = tmp_dir / lm_name
            shutil.copy2(lm_template, lm_out)
            _apply_text_to_docx(lm_out, new_lm)

            email_subject, email_body = _parse_email_output(email_raw)
            eml_bytes = build_eml(email_subject, email_body, cv_out, lm_out)

            zf.write(cv_out, f"{slug}/{cv_name}")
            zf.write(lm_out, f"{slug}/{lm_name}")
            zf.writestr(f"{slug}/email.eml", eml_bytes)

    return str(zip_path), total_tokens